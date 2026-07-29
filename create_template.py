"""
Builds `template.docx` — the Jinja2-tagged Word template the app renders into.

The layout is a faithful recreation of the reference invoice: greyscale, a
two-column header over a heavy rule, rounded grey info panels, a products table
with a dark header band and horizontal rules only, and a right-aligned order
total. Corners are rounded via VML shapes, since Word table cells are always
square.

Run once after cloning:

    python create_template.py

You can open template.docx in Word afterwards and restyle it freely; just keep
the {{ tags }} and the {%tr for item in items %} loop rows intact.
"""

import io
import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor, Inches
from docx.text.paragraph import Paragraph

OUTPUT = "template.docx"

# Masthead logo. Embedded straight into the template — it never changes per
# invoice, so it needs no Jinja tag. If the file is missing the build falls
# back to the {{ company_name }} text so the template is still usable.
LOGO_FILE = "logo.jpeg"
LOGO_WIDTH = Inches(2.7)

# Greyscale palette taken from the reference invoice.
INK = RGBColor(0x11, 0x11, 0x11)      # headings / primary text
BODY = RGBColor(0x2B, 0x2B, 0x2B)     # table body text
MUTED = RGBColor(0x73, 0x73, 0x73)    # labels, address, meta
PAPER = RGBColor(0xFF, 0xFF, 0xFF)    # text on the dark header band

CARD_FILL = "F5F5F5"                  # info card background
HEAD_FILL = "222222"                  # products table header band
RULE = "DCDCDC"                       # hairline row rules
STRONG = "222222"                     # heavy structural rules

HAIRLINE = 4      # 0.5pt — row separators
MEDIUM = 8        # 1.0pt — card edges
HEAVY = 16        # 2.0pt — header underline, totals rule

CELL_PAD = 115    # products table cell padding, in twentieths of a point
CARD_WIDTH = 238  # rounded card width, in points (3.35in column, less a hair)

CONTENT_WIDTH = Inches(7.0)


# --------------------------------------------------------------------------- #
#  Low-level docx helpers
# --------------------------------------------------------------------------- #
# w:tcPr children must appear in this order or Word renders the cell oddly —
# shading and borders in particular come out with hairline gaps at cell joins.
_TC_ORDER = [
    qn(t) for t in (
        "w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders",
        "w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
        "w:vAlign", "w:hideMark",
    )
]


def set_tc_property(cell, element):
    """Replace a w:tcPr child, keeping the schema's required element order."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(element.tag):
        tc_pr.remove(existing)

    position = _TC_ORDER.index(element.tag)
    for child in tc_pr:
        if child.tag in _TC_ORDER and _TC_ORDER.index(child.tag) > position:
            child.addprevious(element)
            return
    tc_pr.append(element)


def shade(cell, hex_fill):
    """Solid background fill for a table cell."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    set_tc_property(cell, el)


def set_borders(cell, *, top=None, bottom=None, left=None, right=None, size=4):
    """
    Set individual cell borders. Pass a hex colour to draw an edge, or None to
    explicitly clear it — which is how the table gets horizontal rules only.
    """
    borders = OxmlElement("w:tcBorders")
    for edge, colour in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        if colour is None:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), colour)
        borders.append(el)
    set_tc_property(cell, borders)


def clear_borders(cell):
    set_borders(cell)


# --------------------------------------------------------------------------- #
#  Rounded boxes
#
#  Word table cells are always square-cornered, so the info cards and status
#  pills are VML rounded rectangles (v:roundrect) with the content living in the
#  shape's text box. docxtpl renders tags inside w:txbxContent like any other
#  text, so the {{ placeholders }} still work.
# --------------------------------------------------------------------------- #
VML_NS = 'xmlns:v="urn:schemas-microsoft-com:vml"'

ARC_CARD = "6554f"    # ~10% corner radius — a soft card


class BoxContent:
    """
    A stand-in container that quacks like a table cell.

    It collects paragraphs so the same para()/run() helpers used elsewhere can
    build the inside of a shape, which is then moved into w:txbxContent.
    """

    def __init__(self):
        self._paragraphs = [self._blank()]

    @staticmethod
    def _blank():
        return Paragraph(parse_xml(f"<w:p {nsdecls('w')}/>"), None)

    @property
    def paragraphs(self):
        return self._paragraphs

    def add_paragraph(self):
        p = self._blank()
        self._paragraphs.append(p)
        return p


def _roundrect(width_pt, height_pt, *, fill, stroke, stroke_pt, arc, inset, fit=True):
    """A <w:p> holding a single rounded rectangle shape with an empty text box."""
    stroke_attr = (
        f'strokecolor="#{stroke}" strokeweight="{stroke_pt}pt"' if stroke else 'stroked="f"'
    )
    # mso-fit-shape-to-text lets the panel grow with its contents, so a long
    # address can never be clipped the way a fixed height would clip it.
    autofit = ";mso-fit-shape-to-text:t" if fit else ""
    return parse_xml(
        f"<w:p {nsdecls('w')} {VML_NS}>"
        f'<w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f"<w:r><w:pict>"
        f'<v:roundrect style="width:{width_pt}pt;height:{height_pt}pt;v-text-anchor:top{autofit}" '
        f'arcsize="{arc}" fillcolor="#{fill}" {stroke_attr}>'
        f'<v:textbox inset="{inset}" style="mso-fit-shape-to-text:t"><w:txbxContent/></v:textbox>'
        f"</v:roundrect>"
        f"</w:pict></w:r></w:p>"
    )


def _place(shape_p, doc):
    """Drop a shape paragraph into the document body."""
    placeholder = doc.add_paragraph()
    placeholder._p.addprevious(shape_p)
    placeholder._p.getparent().remove(placeholder._p)


def rounded_panel(doc, build_left, build_right, *, min_height=44):
    """
    A full-width rounded panel holding two columns of content.

    Both columns live inside one shape rather than in two side-by-side shapes:
    separate shapes would need matching fixed heights, and any content that
    outgrew them would be clipped. A table inside a text box is legal in Word
    (unlike a text box inside a text box, which corrupts the file).
    """
    scratch = Document()
    inner = scratch.add_table(rows=1, cols=3)
    fixed_width(inner, [Inches(3.15), Inches(0.35), Inches(3.15)])
    borderless(inner)
    for cell in inner.rows[0].cells:
        pad(cell, 0)

    build_left(inner.cell(0, 0))
    build_right(inner.cell(0, 2))

    shape_p = _roundrect(
        504, min_height,
        fill=CARD_FILL, stroke=RULE, stroke_pt=1, arc=ARC_CARD,
        inset="12pt,10pt,12pt,10pt",
    )
    txbx = shape_p.find(f".//{qn('w:txbxContent')}")
    txbx.append(inner._tbl)
    # A text box's content must end with a paragraph, never a table.
    txbx.append(parse_xml(f"<w:p {nsdecls('w')}><w:pPr><w:spacing w:after=\"0\" "
                          f'w:line="120" w:lineRule="exact"/></w:pPr></w:p>'))

    _place(shape_p, doc)


def pad(cell, twips):
    """Inner margins for a cell, in twentieths of a point."""
    margins = OxmlElement("w:tcMar")
    for edge, value in (
        ("top", twips), ("bottom", twips),
        ("left", int(twips * 1.1)), ("right", int(twips * 1.1)),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    set_tc_property(cell, margins)


def run(paragraph, text, *, size=10, bold=False, color=INK, caps=False, spacing=None):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.name = "Calibri"
    if caps:
        r.font.all_caps = True
    if spacing:
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(spacing))
        r._element.get_or_add_rPr().append(el)
    return r


def para(container, *, space_before=0, space_after=0, align=None, first=False):
    """Fetch the first paragraph of a cell, or append a new one."""
    p = container.paragraphs[0] if first else container.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def fixed_width(table, widths):
    """
    Lock column widths. Word ignores cell.width unless autofit is off *and* the
    layout is declared fixed; without zeroed cell spacing it also leaves hairline
    white gaps at cell joins, which break up shaded bands and full-width rules.
    """
    table.autofit = False

    tbl_pr = table._tbl.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    spacing = OxmlElement("w:tblCellSpacing")
    spacing.set(qn("w:w"), "0")
    spacing.set(qn("w:type"), "dxa")
    tbl_pr.append(spacing)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]


def borderless(table):
    for row in table.rows:
        for cell in row.cells:
            clear_borders(cell)


def rule(doc, *, colour=STRONG, size=HEAVY, space_before=0, space_after=0):
    """
    A full-width horizontal rule drawn as a paragraph border.

    A vertically merged cell won't render its bottom border reliably, so the
    masthead rule is drawn this way instead — it always spans the text column.
    """
    p = doc.add_paragraph()

    # w:pBdr must precede w:spacing, so append the border before touching
    # paragraph_format (which is what creates the w:spacing element).
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    borders.append(bottom)
    p_pr.append(borders)

    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run().font.size = Pt(2)  # keeps the rule's own line height minimal
    return p


# --------------------------------------------------------------------------- #
#  Building blocks
# --------------------------------------------------------------------------- #
def trimmed_logo(path):
    """
    The supplied logo has a wide white margin baked into the JPEG, which would
    push the mark away from the page margin and leave the masthead looking
    indented. Crop to the artwork so it sits flush left.

    Returns a file-like object, or the original path if Pillow isn't installed.
    """
    try:
        from PIL import Image, ImageChops
    except ImportError:
        return path

    image = Image.open(path).convert("RGB")
    white = Image.new("RGB", image.size, (255, 255, 255))
    mask = ImageChops.difference(image, white).convert("L").point(
        lambda p: 255 if p > 12 else 0
    )
    box = mask.getbbox()
    if not box:
        return path

    buffer = io.BytesIO()
    image.crop(box).save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def card_label(cell, text):
    """Small uppercase letterspaced label — the eyebrow above each card value."""
    p = para(cell, first=True, space_after=4)
    run(p, text, size=7.5, bold=True, color=MUTED, caps=True, spacing=18)


def status_value(box, tag):
    p = para(box)
    run(p, tag, size=12, bold=True, color=INK)




# --------------------------------------------------------------------------- #
#  Template
# --------------------------------------------------------------------------- #
def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    # The stock Normal style carries 8pt space-after and 1.08 line spacing, which
    # silently inflates content inside the fixed-height cards until it clips.
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ------------------------------------------------------------------ header
    header = doc.add_table(rows=3, cols=2)
    fixed_width(header, [Inches(4.0), Inches(3.0)])
    borderless(header)

    # Company name is vertically centred against the stacked invoice block.
    masthead = header.cell(0, 0).merge(header.cell(2, 0))
    masthead.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title = para(masthead, first=True)

    logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), LOGO_FILE)
    if os.path.exists(logo):
        title.add_run().add_picture(trimmed_logo(logo), width=LOGO_WIDTH)
    else:
        run(title, "{{ company_name }}", size=21, bold=True, color=INK)

    kind = para(header.cell(0, 1), first=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(kind, "Invoice", size=10, bold=True, color=MUTED, caps=True, spacing=60)

    number = para(header.cell(1, 1), first=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(number, "{{ invoice_number }}", size=17, bold=True, color=INK)

    date = para(header.cell(2, 1), first=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(date, "{{ invoice_date }}", size=9.5, color=MUTED)

    # Heavy rule anchoring the masthead.
    rule(doc, space_before=2, space_after=14)

    # ------------------------------------------------------- reseller / customer
    def reseller(box):
        card_label(box, "Reseller")
        p = para(box)
        run(p, "{{ reseller_info }}", size=12, bold=True, color=INK)

    def customer(box):
        card_label(box, "Customer / Receiver")
        p = para(box)
        run(p, "{{ customer_name }}", size=12, bold=True, color=INK)
        a = para(box, space_before=3)
        run(a, "{{ customer_address }}", size=9.5, color=MUTED)

    rounded_panel(doc, reseller, customer, min_height=64)
    para(doc, space_after=5)

    # ---------------------------------------------------------------- statuses
    def payment(box):
        card_label(box, "Payment Status")
        status_value(box, "{{ payment_status }}")

    def shipping(box):
        card_label(box, "Shipping Status")
        status_value(box, "{{ shipping_status }}")

    rounded_panel(doc, payment, shipping, min_height=40)
    para(doc, space_after=9)

    # ---------------------------------------------------------------- products
    heading = para(doc, space_after=6)
    run(heading, "Products", size=11, bold=True, color=INK, caps=True, spacing=30)

    headers = ["CODE", "PRODUCT", "QTY", "UNIT PRICE", "LINE TOTAL"]
    widths = [Inches(1.05), Inches(2.95), Inches(0.7), Inches(1.15), Inches(1.15)]

    # Four rows: header, loop-open, the repeating data row, loop-close.
    #
    # docxtpl's row tag must be written as "{%tr ... %}" — NO space after "{%",
    # exactly one space after "tr" — and it must sit in a row of its own, because
    # docxtpl deletes the entire <w:tr> that contains it. Writing "{% tr ... %}"
    # leaves the tag in the XML and Jinja then fails with "unknown tag 'tr'".
    table = doc.add_table(rows=4, cols=5)
    fixed_width(table, widths)

    # Dark header band — the strongest horizontal element on the page.
    #
    # It is one merged cell rather than five: abutting shaded cells leave
    # hairline seams that show as vertical lines across the band, and painting
    # the cell edges to hide them just turns the seams into visible bars.
    # Labels are positioned with tab stops derived from the column widths so
    # they stay aligned with the data below.
    band = table.cell(0, 0).merge(table.cell(0, 4))
    shade(band, HEAD_FILL)
    set_borders(band, top=STRONG, bottom=STRONG, size=MEDIUM)
    pad(band, CELL_PAD)

    inset = Inches(CELL_PAD * 1.1 / 1440)
    edges = [sum(widths[: i + 1], Inches(0)) for i in range(len(widths))]

    p = para(band, first=True)
    stops = p.paragraph_format.tab_stops
    stops.add_tab_stop(widths[0], WD_TAB_ALIGNMENT.LEFT)             # PRODUCT
    for edge in edges[2:]:                                            # right-aligned
        stops.add_tab_stop(edge - inset * 2, WD_TAB_ALIGNMENT.RIGHT)

    for i, text in enumerate(headers):
        if i:
            run(p, "\t")
        run(p, text, size=8, bold=True, color=PAPER, caps=True, spacing=20)

    run(para(table.cell(1, 0), first=True), "{%tr for item in items %}", size=8)

    body = [
        "{{ item.code }}",
        "{{ item.product }}",
        "{{ item.qty }}",
        "{{ item.unit_price }}",
        "{{ item.line_total }}",
    ]
    for i, text in enumerate(body):
        cell = table.cell(2, i)
        set_borders(cell, bottom=RULE, size=HAIRLINE)  # horizontal rules only
        pad(cell, CELL_PAD)
        p = para(cell, first=True, align=WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else None)
        # Codes and money read as the anchors of each row, so they carry weight.
        run(p, text, size=10, bold=i in (0, 4), color=INK if i in (0, 4) else BODY)

    run(para(table.cell(3, 0), first=True), "{%tr endfor %}", size=8)

    for i in (1, 3):
        for cell in table.rows[i].cells:
            clear_borders(cell)

    # ------------------------------------------------------------------- total
    para(doc, space_after=2)

    # A single cell with a right tab stop, rather than two cells — abutting cells
    # leave a hairline gap that would show as a break in the heavy rule above.
    totals = doc.add_table(rows=1, cols=1)
    totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
    fixed_width(totals, [Inches(3.4)])

    cell = totals.cell(0, 0)
    set_borders(cell, top=STRONG, size=HEAVY)
    pad(cell, 110)

    p = para(cell, first=True)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(3.2), WD_TAB_ALIGNMENT.RIGHT)
    run(p, "Order Total", size=10, bold=True, color=MUTED, caps=True, spacing=20)
    run(p, "\t")
    run(p, "{{ order_total }}", size=15, bold=True, color=INK)

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
